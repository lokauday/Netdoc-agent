import os
import json
import streamlit as st
from openai import OpenAI
from utils.parser import parse_config
from docx import Document

# -----------------------------
# Load API key
# -----------------------------
api_key = st.secrets.get("OPENAI_API_KEY") or os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=api_key)

# -----------------------------
# Page setup
# -----------------------------
st.set_page_config(
    page_title="NetDoc AI",
    layout="wide",
    page_icon="🛠️"
)

# -----------------------------
# Global CSS (Hybrid Light UI)
# -----------------------------
st.markdown("""
<style>
/* Top navigation bar */
.navbar {
    background: #f6f9fc;
    padding: 12px 20px;
    border-bottom: 1px solid #dce2e9;
    font-size: 22px;
    font-weight: 600;
    color: #2b4c7e;
}

/* Sidebar styling */
section[data-testid="stSidebar"] {
    background-color: #ffffff !important;
    border-right: 1px solid #e6e9ef;
}

.sidebar-title {
    font-size: 20px;
    font-weight: 700;
    color: #2b4c7e;
}

.block {
    background: #ffffff;
    padding: 18px;
    border-radius: 12px;
    box-shadow: 0 2px 10px rgba(0,0,0,0.06);
    margin-top: 10px;
}

/* Headers */
.section-header {
    font-size: 24px;
    font-weight: 700;
    color: #2b4c7e;
    margin-bottom: 10px;
}
</style>
""", unsafe_allow_html=True)

# -----------------------------
# Top Navigation Bar
# -----------------------------
st.markdown(
    '<div class="navbar">⚡ NetDoc AI — Autonomous Network Documentation Engine</div>',
    unsafe_allow_html=True
)

# -----------------------------
# Sidebar Menu
# -----------------------------
with st.sidebar:
    st.markdown("<div class='sidebar-title'>📁 Menu</div>", unsafe_allow_html=True)

    choice = st.radio(
        "",
        [
            "Dashboard",
            "Upload",
            "Documentation",
            "Audit",
            "Topology",
            "Exports",
            "About"
        ]
    )


# ====================================================
#                   PAGE: Dashboard
# ====================================================
if choice == "Dashboard":
    st.markdown("<div class='section-header'>📊 Dashboard</div>", unsafe_allow_html=True)

    st.write("Welcome to **NetDoc AI** — Upload configs, generate documentation, run audits, and export beautiful reports.")


# ====================================================
#                   PAGE: Upload
# ====================================================
elif choice == "Upload":
    st.markdown("<div class='section-header'>📁 Upload Device Configurations</div>", unsafe_allow_html=True)

    uploaded_files = st.file_uploader(
        "Upload one or more config files",
        type=["txt", "cfg", "log"],
        accept_multiple_files=True
    )

    if uploaded_files and st.button("Generate Report"):
        combined = ""
        for f in uploaded_files:
            combined += f"\n\n# FILE: {f.name}\n"
            combined += f.read().decode("utf-8")

        with st.spinner("Processing configurations..."):
            result = parse_config(combined)

        st.success("Documentation created successfully!")
        st.session_state["report"] = result
        st.session_state["markdown"] = json.dumps(result, indent=2)


# ====================================================
#               PAGE: Documentation
# ====================================================
elif choice == "Documentation":
    st.markdown("<div class='section-header'>📄 Generated Documentation</div>", unsafe_allow_html=True)

    if "report" in st.session_state:
        st.json(st.session_state["report"])
    else:
        st.info("Upload configs to generate documentation.")


# ====================================================
#               PAGE: Security Audit
# ====================================================
elif choice == "Audit":
    st.markdown("<div class='section-header'>🛡 Security Audit</div>", unsafe_allow_html=True)

    if "markdown" not in st.session_state:
        st.info("Upload configs first.")
    else:
        with st.spinner("AI analyzing configuration security…"):
            prompt = f"""
You are a network security auditor.

Review this network configuration JSON and produce:

- Critical issues (HIGH)
- Moderate issues (MEDIUM)
- Best practices (LOW)
- Remediation recommendations

Return clear bullet points.

Config:
{st.session_state["markdown"]}
"""
            ai_out = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}]
            )

        st.subheader("Audit Report")
        st.write(ai_out.choices[0].message["content"])


# ====================================================
#               PAGE: Topology
# ====================================================
elif choice == "Topology":
    st.markdown("<div class='section-header'>🌐 Network Topology Diagram</div>", unsafe_allow_html=True)

    if "markdown" not in st.session_state:
        st.info("Generate documentation first.")
    else:
        with st.spinner("AI generating topology diagram…"):
            prompt = f"""
Create an ASCII topology diagram from this parsed configuration:

{st.session_state["markdown"]}

Rules:
- Show devices, VLANs, subnets, interfaces.
- Clean ASCII. No markdown formatting.
"""
            ai_topo = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}]
            )

        topo_text = ai_topo.choices[0].message["content"]
        st.text(topo_text)
        st.session_state["topology"] = topo_text


# ====================================================
#                   PAGE: Exports
# ====================================================
elif choice == "Exports":
    st.markdown("<div class='section-header'>📤 Export Documentation</div>", unsafe_allow_html=True)

    if "markdown" not in st.session_state:
        st.info("Generate documentation first.")
    else:
        md = st.session_state["markdown"]
        topo = st.session_state.get("topology", "")

        # ---------------- PDF Export ----------------
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter

        def make_pdf():
            import io
            buffer = io.BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            text = c.beginText(40, 750)
            text.setFont("Helvetica", 8)

            for line in md.split("\n"):
                text.textLine(line)

            if topo:
                text.textLine("\n--- TOPOLOGY ---\n")
                for line in topo.split("\n"):
                    text.textLine(line)

            c.drawText(text)
            c.save()
            return buffer.getvalue()

        pdf_data = make_pdf()

        st.download_button(
            "📄 Download PDF",
            data=pdf_data,
            file_name="NetDoc_Report.pdf",
            mime="application/pdf"
        )

        # ---------------- DOCX Export ----------------
        def make_docx():
            doc = Document()
            doc.add_heading("NetDoc AI Report", level=1)
            doc.add_page_break()

            doc.add_heading("Documentation", level=2)
            doc.add_paragraph(md)

            if topo:
                doc.add_heading("Topology Diagram", level=2)
                doc.add_paragraph(topo)

            buffer = io.BytesIO()
            doc.save(buffer)
            return buffer.getvalue()

        import io
        docx_data = make_docx()

        st.download_button(
            "📘 Download DOCX",
            data=docx_data,
            file_name="NetDoc_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        st.success("Export files ready!")


# ====================================================
#                   PAGE: About
# ====================================================
elif choice == "About":
    st.markdown("<div class='section-header'>ℹ️ About NetDoc AI</div>", unsafe_allow_html=True)
    st.write("""
NetDoc AI automatically generates:

- Device documentation  
- VLAN breakdown  
- Interface tables  
- LLDP/CDP neighbors  
- Routing summary  
- Security audit  
- Topology diagrams  
- PDF & DOCX exports  

Enterprise-grade. Fast. Accurate.
""")


