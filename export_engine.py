# ============================================================
#  EXPORT ENGINE — PDF, DOCX, HTML
# ============================================================

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics

from docx import Document
from io import BytesIO
import json
import html

FONT_PATH = "NotoSans-Regular.ttf"

# Register Unicode font
try:
    pdfmetrics.registerFont(TTFont("Noto", FONT_PATH))
    BASE_FONT = "Noto"
except:
    BASE_FONT = "Helvetica"


# ------------------------------------------------------------
#  PDF GENERATION
# ------------------------------------------------------------
def generate_pdf_report(parsed):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)

    styles = getSampleStyleSheet()
    style = styles["Normal"]
    style.fontName = BASE_FONT

    elements = []
    elements.append(Paragraph("<b>NetDoc AI — Network Report</b>", style))
    elements.append(Spacer(1, 12))

    pretty = html.escape(json.dumps(parsed, indent=4)).replace("\n", "<br/>")

    elements.append(Paragraph("<b>Parsed Configuration:</b>", style))
    elements.append(Paragraph(f"<pre>{pretty}</pre>", style))

    doc.build(elements)
    pdf_bytes = buffer.getvalue()
    buffer.close()

    return pdf_bytes


# ------------------------------------------------------------
#  DOCX GENERATION
# ------------------------------------------------------------
def generate_docx_report(parsed):
    doc = Document()

    doc.add_heading("NetDoc AI — Network Report", level=1)
    doc.add_paragraph(json.dumps(parsed, indent=4))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ------------------------------------------------------------
#  HTML GENERATION
# ------------------------------------------------------------
def generate_html_report(parsed):
    pretty = html.escape(json.dumps(parsed, indent=4))
    pretty = pretty.replace(" ", "&nbsp;").replace("\n", "<br>")

    html_data = f"""
        <html>
        <body>
            <h1>NetDoc AI Report</h1>
            <pre>{pretty}</pre>
        </body>
        </html>
    """

    return html_data.encode()


# ------------------------------------------------------------
#  MASTER EXPORT WRAPPER
# ------------------------------------------------------------
def export_all_formats(parsed):
    pdf = generate_pdf_report(parsed)
    docx = generate_docx_report(parsed)
    html = generate_html_report(parsed)

    return pdf, docx, html
