# ===============================================================
#  NetDoc AI — Export Engine (PDF, DOCX, ZIP)
# ===============================================================

import io
import json
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document
from zipfile import ZipFile


# ---------------------------------------------------------------
# PDF EXPORT
# ---------------------------------------------------------------
def export_pdf(audit: dict, topology: str) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)

    styles = getSampleStyleSheet()
    flow = []

    flow.append(Paragraph("<b>NetDoc AI — Security Audit Report</b>", styles["Title"]))
    flow.append(Spacer(1, 12))

    flow.append(Paragraph("<b>Issues:</b>", styles["Heading2"]))
    for i in audit["issues"]:
        flow.append(Paragraph(f"• {i}", styles["Normal"]))

    flow.append(Spacer(1, 10))

    flow.append(Paragraph("<b>Warnings:</b>", styles["Heading2"]))
    for w in audit["warnings"]:
        flow.append(Paragraph(f"• {w}", styles["Normal"]))

    flow.append(Spacer(1, 10))

    flow.append(Paragraph("<b>Info:</b>", styles["Heading2"]))
    for x in audit["info"]:
        flow.append(Paragraph(f"• {x}", styles["Normal"]))

    flow.append(Spacer(1, 10))

    flow.append(Paragraph("<b>Topology Diagram (Mermaid Source)</b>", styles["Heading2"]))
    flow.append(Paragraph(f"<pre>{topology}</pre>", styles["Code"]))

    doc.build(flow)
    return buffer.getvalue()


# ---------------------------------------------------------------
# DOCX EXPORT
# ---------------------------------------------------------------
def export_docx(audit: dict, topology: str) -> bytes:
    doc = Document()

    doc.add_heading("NetDoc AI — Security Audit Report", level=1)

    doc.add_heading("Issues", level=2)
    for i in audit["issues"]:
        doc.add_paragraph(i, style="List Bullet")

    doc.add_heading("Warnings", level=2)
    for w in audit["warnings"]:
        doc.add_paragraph(w, style="List Bullet")

    doc.add_heading("Info", level=2)
    for x in audit["info"]:
        doc.add_paragraph(x, style="List Bullet")

    doc.add_heading("Topology Diagram (Mermaid Source)", level=2)
    doc.add_paragraph(topology)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------
# ZIP EXPORT
# ---------------------------------------------------------------
def export_zip(audit: dict, topology: str, raw_config: str) -> bytes:
    buffer = io.BytesIO()

    with ZipFile(buffer, mode="w") as zipf:
        # JSON
        zipf.writestr("audit.json", json.dumps(audit, indent=4))

        # TOPOLOGY
        zipf.writestr("topology.mmd", topology)

        # RAW CONFIG
        zipf.writestr("config_raw.txt", raw_config)

        # Markdown
        md = (
            "# NetDoc AI — Audit Report\n\n"
            "## Issues\n" + "\n".join(f"- {i}" for i in audit["issues"]) +
            "\n\n## Warnings\n" + "\n".join(f"- {w}" for w in audit["warnings"]) +
            "\n\n## Info\n" + "\n".join(f"- {x}" for x in audit["info"]) +
            "\n\n## Topology\n```mermaid\n" +
            topology +
            "\n```\n"
        )
        zipf.writestr("audit.md", md)

    return buffer.getvalue()
