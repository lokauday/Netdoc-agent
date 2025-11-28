import json
from docx import Document
from reportlab.pdfgen import canvas
from io import BytesIO

# =====================================================
# SECURITY AUDIT ENGINE
# =====================================================

def run_security_audit(report):
    audit = {
        "weak_passwords": [],
        "missing_enable_secret": False,
        "no_bpdu_guard": [],
        "open_vlans": [],
        "stp_issues": [],
        "acl_warnings": [],
    }

    # ---------- Weak Passwords ----------
    text = json.dumps(report)

    weak_keywords = ["password 0", "username admin password", "enable password"]

    for key in weak_keywords:
        if key in text.lower():
            audit["weak_passwords"].append(f"Found weak credential: '{key}'")

    if "enable secret" not in text.lower():
        audit["missing_enable_secret"] = True

    # ---------- BPDU Guard ----------
    if "interfaces" in report:
        for iface in report["interfaces"]:
            if "Gig" in iface.get("name", "") and "bpduguard enable" not in text.lower():
                audit["no_bpdu_guard"].append(iface["name"])

    # ---------- VLAN Issues ----------
    if "vlans" in report:
        for vlan in report["vlans"]:
            if len(vlan.get("ports", [])) == 0:
                audit["open_vlans"].append(f"VLAN {vlan.get('vlan_id')} has no active ports")

    # ---------- STP ----------
    if "spanning-tree mode pvst" not in text.lower() and "spanning-tree mode rapid-pvst" not in text.lower():
        audit["stp_issues"].append("STP mode missing (PVST or Rapid-PVST recommended)")

    # ---------- ACL ----------
    if "access-list" not in text.lower():
        audit["acl_warnings"].append("No access-lists found — device may be unprotected")

    return audit


# =====================================================
# TOPOLOGY DIAGRAM (MERMAID)
# =====================================================

def generate_topology_mermaid(report):
    """
    Creates a simple but beautiful Mermaid topology diagram.
    """

    mermaid = ["graph TD;"]

    if "neighbors" not in report:
        return "graph TD;\nA[No Neighbors Found];"

    for n in report["neighbors"]:
        local = n.get("local_interface", "Unknown")
        remote = n.get("neighbor_device", "Device")
        r_int = n.get("neighbor_interface", "Port")

        mermaid.append(f'{local} -- "{r_int}" --> {remote}')

    return "\n".join(mermaid)


# ============================================================
#  EXPORT ENGINE — PDF + DOCX + HTML (Streamlit Safe)
# ============================================================

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from io import BytesIO
from docx import Document


# ------------------------------------------------------------
# 1) PDF EXPORT
# ------------------------------------------------------------
def export_pdf(report_dict):
    buffer = BytesIO()

    # PDF Canvas
    pdf = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    pdf.setFont("Helvetica", 10)

    y = height - 60

    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(50, y, "NetDoc AI — Network Documentation Report")
    y -= 30

    pdf.setFont("Helvetica", 10)

    def write_line(text, step=14):
        nonlocal y
        if y < 50:
            pdf.showPage()
            pdf.setFont("Helvetica", 10)
            y = height - 50
        pdf.drawString(50, y, text)
        y -= step

    write_line("")

    # Dump dictionary to PDF
    import json
    text = json.dumps(report_dict, indent=2)

    for line in text.split("\n"):
        write_line(line)

    pdf.save()
    pdf_bytes = buffer.getvalue()
    buffer.close()

    return pdf_bytes


# ------------------------------------------------------------
# 2) DOCX EXPORT
# ------------------------------------------------------------
def export_docx(report_dict):
    doc = Document()

    doc.add_heading("NetDoc AI — Network Documentation Report", level=1)

    import json
    text = json.dumps(report_dict, indent=2)

    doc.add_paragraph(text)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


# ------------------------------------------------------------
# 3) HTML EXPORT
# ------------------------------------------------------------
def export_html(report_dict):
    html = """
    <html>
    <head>
        <title>NetDoc AI Report</title>
        <style>
            body {
                font-family: Arial;
                padding: 25px;
                background: #f7f7f7;
            }
            pre {
                background: #fff;
                padding: 15px;
                border-radius: 8px;
                border: 1px solid #ddd;
                font-size: 14px;
            }
        </style>
    </head>
    <body>
        <h2>NetDoc AI — Network Documentation Report</h2>
        <pre>{}</pre>
    </body>
    </html>
    """.format(report_dict)

    return html.encode("utf-8")


# ------------------------------------------------------------
# 4) MASTER EXPORT FUNCTION
# ------------------------------------------------------------
def export_all_formats(report_dict):
    pdf_bytes = export_pdf(report_dict)
    docx_bytes = export_docx(report_dict)
    html_bytes = export_html(report_dict)

    return pdf_bytes, docx_bytes, html_bytes
